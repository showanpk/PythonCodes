from pathlib import Path
from docx import Document
from faster_whisper import WhisperModel
import sys


def format_timestamp(seconds: float) -> str:
    total_ms = int(seconds * 1000)
    hours = total_ms // 3_600_000
    minutes = (total_ms % 3_600_000) // 60_000
    secs = (total_ms % 60_000) // 1000
    ms = total_ms % 1000
    return f"{hours:02}:{minutes:02}:{secs:02}.{ms:03}"


def transcribe_ogg_to_docx(
    input_file: str,
    output_docx: str = None,
    model_size: str = "small",
    language: str = None,
    device: str = "cpu",
    compute_type: str = "int8",
    include_timestamps: bool = True,
) -> Path:
    input_path = Path(input_file)

    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    if output_docx is None:
        output_path = input_path.with_suffix(".docx")
    else:
        candidate = Path(output_docx).expanduser()

        # If a directory is provided, write a .docx with the input stem inside it.
        if candidate.exists() and candidate.is_dir():
            output_path = candidate / f"{input_path.stem}.docx"
        elif candidate.suffix.lower() != ".docx":
            output_path = candidate / f"{input_path.stem}.docx"
        else:
            output_path = candidate

    output_path.parent.mkdir(parents=True, exist_ok=True)

    print("Loading model...")
    model = WhisperModel(model_size, device=device, compute_type=compute_type)

    print("Transcribing audio...")
    segments, info = model.transcribe(
        str(input_path),
        beam_size=5,
        language=language,   # set to "en" if you want English only
        vad_filter=True
    )

    # Important: segments is a generator, so force evaluation
    segments = list(segments)

    full_text_parts = []
    timestamped_lines = []

    for seg in segments:
        text = seg.text.strip()
        if not text:
            continue

        full_text_parts.append(text)

        if include_timestamps:
            start_ts = format_timestamp(seg.start)
            end_ts = format_timestamp(seg.end)
            timestamped_lines.append(f"[{start_ts} - {end_ts}] {text}")

    full_text = "\n".join(full_text_parts).strip()

    print("Creating Word document...")
    doc = Document()
    doc.add_heading("Audio Transcript", level=1)

    doc.add_paragraph(f"Source file: {input_path.name}")
    doc.add_paragraph(f"Detected language: {getattr(info, 'language', 'unknown')}")
    doc.add_paragraph(
        f"Language probability: "
        f"{getattr(info, 'language_probability', 0):.4f}"
        if hasattr(info, "language_probability")
        else "Language probability: unknown"
    )

    doc.add_heading("Clean Transcript", level=2)
    if full_text:
        for para in full_text.split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    else:
        doc.add_paragraph("No speech detected.")

    if include_timestamps:
        doc.add_heading("Transcript With Timestamps", level=2)
        if timestamped_lines:
            for line in timestamped_lines:
                doc.add_paragraph(line)
        else:
            doc.add_paragraph("No timestamped transcript available.")

    doc.save(output_path)
    print(f"Done. Saved Word file to: {output_path.resolve()}")
    return output_path


if __name__ == "__main__":
    """
    Usage:
        python ogg_to_word.py "C:\\path\\to\\audio.ogg"
        python ogg_to_word.py "C:\\path\\to\\audio.ogg" "C:\\path\\to\\output.docx"

    Optional quick edits inside the function call below:
    - model_size="small"   -> faster
    - model_size="medium"  -> usually better accuracy, slower
    - language="en"        -> force English
    """

    def _clean_path(raw: str) -> str:
        return raw.strip().strip('"').strip("'")

    default_input = Path.home() / "Downloads" / "WhatsApp Ptt 2026-03-26 at 12.37.03.ogg"
    default_output = Path.home() / "Downloads"

    if len(sys.argv) >= 2:
        input_audio = _clean_path(sys.argv[1])
        output_docx = _clean_path(sys.argv[2]) if len(sys.argv) > 2 else None
    elif default_input.exists():
        input_audio = str(default_input)
        output_docx = str(default_output)
        print(f"Using default input path: {input_audio}")
        print(f"Using default output path: {output_docx}")
    else:
        print("No input path was provided in the command.")
        print("Paste a full audio path, e.g. C:\\Users\\shonk\\Downloads\\file.ogg")

        input_audio = _clean_path(input("Input .ogg path: "))
        if not input_audio:
            print("Input path is required.")
            sys.exit(1)

        output_raw = _clean_path(input("Output .docx path (optional, press Enter to auto-create): "))
        output_docx = output_raw if output_raw else None

    transcribe_ogg_to_docx(
        input_file=input_audio,
        output_docx=output_docx,
        model_size="small",
        language=None,      # change to "en" for English only
        device="cpu",       # change to "cuda" if you have a supported GPU
        compute_type="int8",
        include_timestamps=True,
    )