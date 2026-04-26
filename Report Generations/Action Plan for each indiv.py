import re
import os
from pathlib import Path
from datetime import datetime

import pyodbc
from docx import Document

# =========================================================
# CONFIG
# =========================================================

TEMPLATE_PATH = r"C:\Users\shonk\Saheli Hub\Saheli Hub - Forms\Showan\Rakhya\5 Ways\Action Plan.docx"
OUTPUT_DIR = r"C:\Users\shonk\Saheli Hub\Saheli Hub - Forms\Showan\Rakhya\5 Ways\Generated"

SQL_SERVER = r"20.68.160.100,1433"
SQL_DATABASE = "SahelihubCRM"
SQL_USERNAME = "saheli_app"
SQL_PASSWORD = "309183"
USE_WINDOWS_AUTH = False

FUNDING_PROJECT_ID = None
# Example:
# FUNDING_PROJECT_ID = "851ec328-c301-450d-82f0-dda93f52407e"

ASSESSMENT_DATE_FROM = "2025-08-01"

DEBUG = False

# =========================================================
# DB
# =========================================================

def env_bool(name: str, default: bool) -> bool:
    val = os.getenv(name)
    if val is None:
        return default
    return val.strip().lower() in {"1", "true", "yes", "y", "on"}


def resolve_db_config() -> dict:
    return {
        "server": os.getenv("DB_SERVER", SQL_SERVER),
        "database": os.getenv("DB_DATABASE", SQL_DATABASE),
        "uid": os.getenv("DB_USERNAME", SQL_USERNAME),
        "pwd": os.getenv("DB_PASSWORD", SQL_PASSWORD),
        "use_windows_auth": env_bool("DB_TRUSTED_CONNECTION", USE_WINDOWS_AUTH),
    }


def get_connection_string() -> str:
    db = resolve_db_config()
    if db["use_windows_auth"]:
        return (
            f"DRIVER={{ODBC Driver 17 for SQL Server}};"
            f"SERVER={db['server']};"
            f"DATABASE={db['database']};"
            f"Trusted_Connection=yes;"
            "TrustServerCertificate=yes;"
        )
    return (
        f"DRIVER={{ODBC Driver 17 for SQL Server}};"
        f"SERVER={db['server']};"
        f"DATABASE={db['database']};"
        f"UID={db['uid']};"
        f"PWD={db['pwd']};"
        "TrustServerCertificate=yes;"
    )


def get_connection():
    try:
        return pyodbc.connect(get_connection_string())
    except pyodbc.Error as ex:
        db = resolve_db_config()
        auth_mode = "Windows Integrated Authentication" if db["use_windows_auth"] else "SQL Authentication"
        raise RuntimeError(
            "Database connection failed using "
            f"{auth_mode}. "
            "Set DB_TRUSTED_CONNECTION=false to force SQL login, "
            "or set DB_USERNAME/DB_PASSWORD with valid credentials."
        ) from ex


def fetch_assessment_pairs(conn) -> list[dict]:
    """
    Return one row per participant with:
    - current/latest assessment
    - previous assessment immediately before the current one (if exists)

    Returns participants with at least one current assessment.
    """
    sql = """
    WITH RankedAssessments AS
    (
        SELECT
            p.ParticipantID,
            p.SaheliCardNumber,
            p.FullName,
            p.Address,
            p.Postcode,
            p.MobileNumber,
            p.Email,
            pfp.FundingProjectId,

            vpa.AssessmentID,
            vpa.AssessmentNumber,
            vpa.AssessmentDate,
            vpa.StaffMember,
            vpa.Site,
            vpa.AimsGoals,
            vpa.AimsDescription,
            vpa.Barriers,
            vpa.BarrierComments,
            vpa.WeightKg,
            vpa.HeightCm,
            vpa.Bmivalue,
            vpa.Bmicategory,
            vpa.WaistHipRatio,
            vpa.BodyFatScore,
            vpa.VisceralFatScore,
            vpa.ConfidenceToJoin,
            vpa.NumberOfHobbies,
            vpa.CommunityInvolvement,
            vpa.ServiceAwareness,
            vpa.SystolicBp,
            vpa.DiastolicBp,
            vpa.Bplevel,
            vpa.HeartAge,
            vpa.DiabetesType,
            vpa.HbA1c,
            vpa.RiskStratification,
            vpa.Nourishment,
            vpa.Movement,
            vpa.SleepQuality,
            vpa.HappySelf,
            vpa.Resilience,
            vpa.ScreenTime,
            vpa.ActiveDaysPerWeek,
            vpa.ActivityLevel,
            vpa.PreferredActivities,
            vpa.NextReviewDate,
            vpa.LackCompanionship,
            vpa.FeelLeftOut,
            vpa.FeelIsolated,
            vpa.FeelingOptimistic,
            vpa.FeelingUseful,
            vpa.FeelingRelaxed,
            vpa.FeelingConfident,
            vpa.FeelingCheerful,
            ROW_NUMBER() OVER
            (
                PARTITION BY p.ParticipantID
                ORDER BY vpa.AssessmentDate DESC, vpa.AssessmentID DESC
            ) AS rn
        FROM ParticipantFundingProjects pfp
        INNER JOIN Participants p
            ON pfp.ParticipantID = p.ParticipantID
        INNER JOIN vw_Participants_Assessments vpa
            ON p.SaheliCardNumber = vpa.SaheliCardNumber
        WHERE vpa.AssessmentDate >= ?
          AND (? IS NULL OR pfp.FundingProjectId = ?)
    )
    SELECT
        cur.ParticipantID,
        cur.SaheliCardNumber,
        cur.FullName,
        cur.Address,
        cur.Postcode,
        cur.MobileNumber,
        cur.Email,
        cur.FundingProjectId,

        -- latest/current
        cur.StaffMember               AS AdvisorName,
        cur.AssessmentID              AS CurrentAssessmentID,
        cur.AssessmentNumber          AS CurrentAssessmentNumber,
        cur.AssessmentDate            AS CurrentAssessmentDate,
        cur.Site                      AS CurrentSite,
        cur.AimsGoals                 AS CurrentAimsGoals,
        cur.AimsDescription           AS CurrentAimsDescription,
        cur.Barriers                  AS CurrentBarriers,
        cur.BarrierComments           AS CurrentBarrierComments,
        cur.WeightKg                  AS CurrentWeightKg,
        cur.HeightCm                  AS CurrentHeightCm,
        cur.Bmivalue                  AS CurrentBmivalue,
        cur.Bmicategory               AS CurrentBmicategory,
        cur.WaistHipRatio             AS CurrentWaistHipRatio,
        cur.BodyFatScore              AS CurrentBodyFatScore,
        cur.VisceralFatScore          AS CurrentVisceralFatScore,
        cur.ConfidenceToJoin          AS CurrentConfidenceToJoin,
        cur.NumberOfHobbies           AS CurrentNumberOfHobbies,
        cur.CommunityInvolvement      AS CurrentCommunityInvolvement,
        cur.ServiceAwareness          AS CurrentServiceAwareness,
        cur.SystolicBp                AS CurrentSystolicBp,
        cur.DiastolicBp               AS CurrentDiastolicBp,
        cur.Bplevel                   AS CurrentBplevel,
        cur.HeartAge                  AS CurrentHeartAge,
        cur.DiabetesType              AS CurrentDiabetesType,
        cur.HbA1c                     AS CurrentHbA1c,
        cur.RiskStratification        AS CurrentRiskStratification,
        cur.Nourishment               AS CurrentNourishment,
        cur.Movement                  AS CurrentMovement,
        cur.SleepQuality              AS CurrentSleepQuality,
        cur.HappySelf                 AS CurrentHappySelf,
        cur.Resilience                AS CurrentResilience,
        cur.ScreenTime                AS CurrentScreenTime,
        cur.ActiveDaysPerWeek         AS CurrentActiveDaysPerWeek,
        cur.ActivityLevel             AS CurrentActivityLevel,
        cur.PreferredActivities       AS CurrentPreferredActivities,
        cur.NextReviewDate            AS CurrentNextReviewDate,
        cur.LackCompanionship         AS CurrentLackCompanionship,
        cur.FeelLeftOut               AS CurrentFeelLeftOut,
        cur.FeelIsolated              AS CurrentFeelIsolated,
        cur.FeelingOptimistic         AS CurrentFeelingOptimistic,
        cur.FeelingUseful             AS CurrentFeelingUseful,
        cur.FeelingRelaxed            AS CurrentFeelingRelaxed,
        cur.FeelingConfident          AS CurrentFeelingConfident,
        cur.FeelingCheerful           AS CurrentFeelingCheerful,

        -- previous
        prev.AssessmentID             AS PreviousAssessmentID,
        prev.AssessmentNumber         AS PreviousAssessmentNumber,
        prev.AssessmentDate           AS PreviousAssessmentDate,
        prev.AimsGoals                AS PreviousAimsGoals,
        prev.AimsDescription          AS PreviousAimsDescription,
        prev.Barriers                 AS PreviousBarriers,
        prev.BarrierComments          AS PreviousBarrierComments,
        prev.WeightKg                 AS PreviousWeightKg,
        prev.HeightCm                 AS PreviousHeightCm,
        prev.Bmivalue                 AS PreviousBmivalue,
        prev.Bmicategory              AS PreviousBmicategory,
        prev.WaistHipRatio            AS PreviousWaistHipRatio,
        prev.BodyFatScore             AS PreviousBodyFatScore,
        prev.VisceralFatScore         AS PreviousVisceralFatScore,
        prev.ConfidenceToJoin         AS PreviousConfidenceToJoin,
        prev.NumberOfHobbies          AS PreviousNumberOfHobbies,
        prev.CommunityInvolvement     AS PreviousCommunityInvolvement,
        prev.ServiceAwareness         AS PreviousServiceAwareness,
        prev.SystolicBp               AS PreviousSystolicBp,
        prev.DiastolicBp              AS PreviousDiastolicBp,
        prev.Bplevel                  AS PreviousBplevel,
        prev.HeartAge                 AS PreviousHeartAge,
        prev.DiabetesType             AS PreviousDiabetesType,
        prev.HbA1c                    AS PreviousHbA1c,
        prev.RiskStratification       AS PreviousRiskStratification,
        prev.Nourishment              AS PreviousNourishment,
        prev.Movement                 AS PreviousMovement,
        prev.SleepQuality             AS PreviousSleepQuality,
        prev.HappySelf                AS PreviousHappySelf,
        prev.Resilience               AS PreviousResilience,
        prev.ScreenTime               AS PreviousScreenTime,
        prev.ActiveDaysPerWeek        AS PreviousActiveDaysPerWeek,
        prev.ActivityLevel            AS PreviousActivityLevel,
        prev.PreferredActivities      AS PreviousPreferredActivities,
        prev.NextReviewDate           AS PreviousNextReviewDate,
        prev.LackCompanionship        AS PreviousLackCompanionship,
        prev.FeelLeftOut              AS PreviousFeelLeftOut,
        prev.FeelIsolated             AS PreviousFeelIsolated,
        prev.FeelingOptimistic        AS PreviousFeelingOptimistic,
        prev.FeelingUseful            AS PreviousFeelingUseful,
        prev.FeelingRelaxed           AS PreviousFeelingRelaxed,
        prev.FeelingConfident         AS PreviousFeelingConfident,
        prev.FeelingCheerful          AS PreviousFeelingCheerful
    FROM RankedAssessments cur
    LEFT JOIN RankedAssessments prev
        ON cur.ParticipantID = prev.ParticipantID
       AND prev.rn = 2
    WHERE cur.rn = 1
    ORDER BY cur.SaheliCardNumber;
    """

    cur = conn.cursor()
    rows = cur.execute(sql, ASSESSMENT_DATE_FROM, FUNDING_PROJECT_ID, FUNDING_PROJECT_ID).fetchall()
    columns = [c[0] for c in cur.description]
    return [dict(zip(columns, row)) for row in rows]

# =========================================================
# HELPERS
# =========================================================

def safe_text(value) -> str:
    if value is None:
        return ""
    return str(value).strip()


def fmt_date(value) -> str:
    if not value:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y")
    try:
        return datetime.fromisoformat(str(value)).strftime("%d/%m/%Y")
    except Exception:
        return str(value)


def safe_filename(value: str) -> str:
    value = safe_text(value) or "Unknown"
    return re.sub(r'[\\/*?:"<>|]+', "_", value)


def combine_address(row: dict) -> str:
    parts = [safe_text(row.get("Address")), safe_text(row.get("Postcode"))]
    return ", ".join([p for p in parts if p])


def is_numberish(value) -> bool:
    try:
        if value is None or str(value).strip() == "":
            return False
        float(str(value))
        return True
    except Exception:
        return False


def compare_value(previous, current, label: str, unit: str = "") -> str:
    prev = safe_text(previous)
    curr = safe_text(current)

    if not prev and not curr:
        return ""

    suffix = f" {unit}".rstrip()

    if prev and curr:
        if is_numberish(prev) and is_numberish(curr):
            prev_num = float(prev)
            curr_num = float(curr)
            diff = curr_num - prev_num
            direction = "increased" if diff > 0 else "decreased" if diff < 0 else "remained the same"
            if diff == 0:
                return f"{label}: {curr}{suffix} (no change from {prev}{suffix})"
            return f"{label}: {curr}{suffix} ({direction} from {prev}{suffix})"
        return f"{label}: {curr} (previously {prev})"

    if curr:
        return f"{label}: {curr}{suffix}"

    return ""


def value_only(current, label: str, unit: str = "") -> str:
    curr = safe_text(current)
    if not curr:
        return ""
    suffix = f" {unit}".rstrip()
    return f"{label}: {curr}{suffix}"


def make_lines(items: list[str]) -> str:
    return "\n".join([x for x in items if safe_text(x)]).strip()


def split_goals(row: dict) -> list[str]:
    raw = safe_text(row.get("CurrentAimsGoals"))
    if not raw:
        raw = safe_text(row.get("PreviousAimsGoals"))

    if not raw:
        return ["", "", ""]

    parts = [x.strip(" -•\t\r\n") for x in re.split(r"[;\n\r]+", raw) if x.strip()]
    while len(parts) < 3:
        parts.append("")
    return parts[:3]


def has_previous_assessment(row: dict) -> bool:
    """Check if participant has a previous assessment for comparison."""
    return bool(row.get("PreviousAssessmentID")) and bool(row.get("PreviousAssessmentDate"))


def has_current_assessment(row: dict) -> bool:
    """Check if participant has a current assessment."""
    return bool(row.get("CurrentAssessmentID")) and bool(row.get("CurrentAssessmentDate"))

# =========================================================
# CONTENT BUILDERS
# =========================================================

def build_background_text(row: dict) -> str:
    parts = []

    if safe_text(row.get("CurrentAimsDescription")):
        parts.append(f"Current aims description: {safe_text(row.get('CurrentAimsDescription'))}")

    if safe_text(row.get("CurrentBarriers")):
        barrier_text = safe_text(row.get("CurrentBarriers"))
        if any(word in barrier_text.lower() for word in ["employment", "work", "family", "depend", "care", "education", "training"]):
            parts.append(f"Background considerations: {barrier_text}")

    if safe_text(row.get("CurrentBarrierComments")):
        parts.append(f"Additional notes: {safe_text(row.get('CurrentBarrierComments'))}")

    if not parts:
        parts.append("Background details not available from previous and current assessment records.")

    return "\n\n".join(parts)


def build_health_week1_text(row: dict) -> str:
    """Build previous assessment health data (only if previous assessment exists)."""
    if not has_previous_assessment(row):
        return ""
    
    lines = [
        value_only(row.get("PreviousWeightKg"), "Weight", "kg"),
        value_only(row.get("PreviousHeightCm"), "Height", "cm"),
        value_only(row.get("PreviousBmivalue"), "BMI"),
        value_only(row.get("PreviousBmicategory"), "BMI category"),
        value_only(row.get("PreviousWaistHipRatio"), "Waist hip ratio"),
        value_only(row.get("PreviousBodyFatScore"), "Body fat score"),
        value_only(row.get("PreviousVisceralFatScore"), "Visceral fat score"),
        value_only(
            f"{safe_text(row.get('PreviousSystolicBp'))}/{safe_text(row.get('PreviousDiastolicBp'))}"
            if safe_text(row.get("PreviousSystolicBp")) or safe_text(row.get("PreviousDiastolicBp"))
            else "",
            "Blood pressure"
        ),
        value_only(row.get("PreviousBplevel"), "BP level"),
        value_only(row.get("PreviousHeartAge"), "Heart age"),
        value_only(row.get("PreviousDiabetesType"), "Diabetes type"),
        value_only(row.get("PreviousHbA1c"), "HbA1c"),
        value_only(row.get("PreviousRiskStratification"), "Risk stratification"),
        value_only(row.get("PreviousSleepQuality"), "Sleep quality"),
        value_only(row.get("PreviousMovement"), "Movement"),
        value_only(row.get("PreviousNourishment"), "Nourishment"),
        value_only(row.get("PreviousConfidenceToJoin"), "Confidence to join"),
        value_only(row.get("PreviousActivityLevel"), "Activity level"),
        value_only(row.get("PreviousActiveDaysPerWeek"), "Active days per week"),
        value_only(row.get("PreviousPreferredActivities"), "Preferred activities"),
        value_only(row.get("PreviousHappySelf"), "Happy self"),
        value_only(row.get("PreviousResilience"), "Resilience"),
        value_only(row.get("PreviousScreenTime"), "Screen time"),
    ]
    return make_lines(lines)


def build_health_week12_text(row: dict) -> str:
    """Build current assessment health data with comparisons if previous exists."""
    if has_previous_assessment(row):
        # Show comparisons with previous assessment
        lines = [
            compare_value(row.get("PreviousWeightKg"), row.get("CurrentWeightKg"), "Weight", "kg"),
            compare_value(row.get("PreviousBmivalue"), row.get("CurrentBmivalue"), "BMI"),
            compare_value(row.get("PreviousBmicategory"), row.get("CurrentBmicategory"), "BMI category"),
            compare_value(row.get("PreviousBodyFatScore"), row.get("CurrentBodyFatScore"), "Body fat score"),
            compare_value(row.get("PreviousVisceralFatScore"), row.get("CurrentVisceralFatScore"), "Visceral fat score"),
            compare_value(
                f"{safe_text(row.get('PreviousSystolicBp'))}/{safe_text(row.get('PreviousDiastolicBp'))}"
                if safe_text(row.get("PreviousSystolicBp")) or safe_text(row.get("PreviousDiastolicBp"))
                else "",
                f"{safe_text(row.get('CurrentSystolicBp'))}/{safe_text(row.get('CurrentDiastolicBp'))}"
                if safe_text(row.get("CurrentSystolicBp")) or safe_text(row.get("CurrentDiastolicBp"))
                else "",
                "Blood pressure"
            ),
            compare_value(row.get("PreviousHeartAge"), row.get("CurrentHeartAge"), "Heart age"),
            compare_value(row.get("PreviousDiabetesType"), row.get("CurrentDiabetesType"), "Diabetes type"),
            compare_value(row.get("PreviousHbA1c"), row.get("CurrentHbA1c"), "HbA1c"),
            compare_value(row.get("PreviousRiskStratification"), row.get("CurrentRiskStratification"), "Risk stratification"),
            compare_value(row.get("PreviousSleepQuality"), row.get("CurrentSleepQuality"), "Sleep quality"),
            compare_value(row.get("PreviousMovement"), row.get("CurrentMovement"), "Movement"),
            compare_value(row.get("PreviousNourishment"), row.get("CurrentNourishment"), "Nourishment"),
            compare_value(row.get("PreviousNumberOfHobbies"), row.get("CurrentNumberOfHobbies"), "Number of hobbies"),
            compare_value(row.get("PreviousCommunityInvolvement"), row.get("CurrentCommunityInvolvement"), "Community involvement"),
            compare_value(row.get("PreviousServiceAwareness"), row.get("CurrentServiceAwareness"), "Service awareness"),
            compare_value(row.get("PreviousLackCompanionship"), row.get("CurrentLackCompanionship"), "Lack companionship"),
            compare_value(row.get("PreviousFeelLeftOut"), row.get("CurrentFeelLeftOut"), "Feel left out"),
            compare_value(row.get("PreviousFeelIsolated"), row.get("CurrentFeelIsolated"), "Feel isolated"),
            compare_value(row.get("PreviousFeelingOptimistic"), row.get("CurrentFeelingOptimistic"), "Feeling optimistic"),
            compare_value(row.get("PreviousFeelingUseful"), row.get("CurrentFeelingUseful"), "Feeling useful"),
            compare_value(row.get("PreviousFeelingRelaxed"), row.get("CurrentFeelingRelaxed"), "Feeling relaxed"),
            compare_value(row.get("PreviousFeelingConfident"), row.get("CurrentFeelingConfident"), "Feeling confident"),
            compare_value(row.get("PreviousFeelingCheerful"), row.get("CurrentFeelingCheerful"), "Feeling cheerful"),
            compare_value(fmt_date(row.get("PreviousNextReviewDate")), fmt_date(row.get("CurrentNextReviewDate")), "Next review date"),
        ]
    else:
        # No previous assessment - show current values only
        lines = [
            value_only(row.get("CurrentWeightKg"), "Weight", "kg"),
            value_only(row.get("CurrentHeightCm"), "Height", "cm"),
            value_only(row.get("CurrentBmivalue"), "BMI"),
            value_only(row.get("CurrentBmicategory"), "BMI category"),
            value_only(row.get("CurrentWaistHipRatio"), "Waist hip ratio"),
            value_only(row.get("CurrentBodyFatScore"), "Body fat score"),
            value_only(row.get("CurrentVisceralFatScore"), "Visceral fat score"),
            value_only(
                f"{safe_text(row.get('CurrentSystolicBp'))}/{safe_text(row.get('CurrentDiastolicBp'))}"
                if safe_text(row.get("CurrentSystolicBp")) or safe_text(row.get("CurrentDiastolicBp"))
                else "",
                "Blood pressure"
            ),
            value_only(row.get("CurrentBplevel"), "BP level"),
            value_only(row.get("CurrentHeartAge"), "Heart age"),
            value_only(row.get("CurrentDiabetesType"), "Diabetes type"),
            value_only(row.get("CurrentHbA1c"), "HbA1c"),
            value_only(row.get("CurrentRiskStratification"), "Risk stratification"),
            value_only(row.get("CurrentSleepQuality"), "Sleep quality"),
            value_only(row.get("CurrentMovement"), "Movement"),
            value_only(row.get("CurrentNourishment"), "Nourishment"),
            value_only(row.get("CurrentConfidenceToJoin"), "Confidence to join"),
            value_only(row.get("CurrentActivityLevel"), "Activity level"),
            value_only(row.get("CurrentActiveDaysPerWeek"), "Active days per week"),
            value_only(row.get("CurrentPreferredActivities"), "Preferred activities"),
            value_only(row.get("CurrentHappySelf"), "Happy self"),
            value_only(row.get("CurrentResilience"), "Resilience"),
            value_only(row.get("CurrentScreenTime"), "Screen time"),
        ]
    return make_lines(lines)


def build_support_text(row: dict) -> str:
    lines = [
        f"Advisor / Staff Member: {safe_text(row.get('AdvisorName'))}" if safe_text(row.get("AdvisorName")) else "",
        f"Site: {safe_text(row.get('CurrentSite'))}" if safe_text(row.get("CurrentSite")) else "",
        f"Current barriers: {safe_text(row.get('CurrentBarriers'))}" if safe_text(row.get("CurrentBarriers")) else "",
        f"Barrier comments: {safe_text(row.get('CurrentBarrierComments'))}" if safe_text(row.get("CurrentBarrierComments")) else "",
        f"Service awareness: {safe_text(row.get('CurrentServiceAwareness'))}" if safe_text(row.get("CurrentServiceAwareness")) else "",
        f"Community involvement: {safe_text(row.get('CurrentCommunityInvolvement'))}" if safe_text(row.get("CurrentCommunityInvolvement")) else "",
    ]
    return make_lines(lines)


def build_supported_text(row: dict) -> str:
    lines = [
        compare_value(row.get("PreviousFeelingOptimistic"), row.get("CurrentFeelingOptimistic"), "Feeling optimistic"),
        compare_value(row.get("PreviousFeelingUseful"), row.get("CurrentFeelingUseful"), "Feeling useful"),
        compare_value(row.get("PreviousFeelingRelaxed"), row.get("CurrentFeelingRelaxed"), "Feeling relaxed"),
        compare_value(row.get("PreviousFeelingConfident"), row.get("CurrentFeelingConfident"), "Feeling confident"),
        compare_value(row.get("PreviousFeelingCheerful"), row.get("CurrentFeelingCheerful"), "Feeling cheerful"),
        compare_value(row.get("PreviousLackCompanionship"), row.get("CurrentLackCompanionship"), "Lack companionship"),
        compare_value(row.get("PreviousFeelLeftOut"), row.get("CurrentFeelLeftOut"), "Feeling left out"),
        compare_value(row.get("PreviousFeelIsolated"), row.get("CurrentFeelIsolated"), "Feeling isolated"),
    ]
    return make_lines(lines)


def build_referral_text(row: dict) -> str:
    lines = [
        f"Reason based on current barriers: {safe_text(row.get('CurrentBarriers'))}" if safe_text(row.get("CurrentBarriers")) else "",
        f"Current barrier comments: {safe_text(row.get('CurrentBarrierComments'))}" if safe_text(row.get("CurrentBarrierComments")) else "",
        f"Current risk stratification: {safe_text(row.get('CurrentRiskStratification'))}" if safe_text(row.get("CurrentRiskStratification")) else "",
    ]
    return make_lines(lines)


def build_comments_text(row: dict) -> str:
    # Assessment metadata removed - user does not need these details
    return ""


def build_goal_review_text(goal_text: str, row: dict) -> str:
    if not safe_text(goal_text):
        return ""

    notes = []
    if safe_text(row.get("CurrentAimsDescription")):
        notes.append(f"Current review note: {safe_text(row.get('CurrentAimsDescription'))}")
    if safe_text(row.get("CurrentBarriers")):
        notes.append(f"Barriers affecting progress: {safe_text(row.get('CurrentBarriers'))}")
    if not notes:
        notes.append("To be reviewed at week 12.")
    return "\n".join(notes)

# =========================================================
# WORD TABLE MAPPING - SINGLE TABLE TEMPLATE
# =========================================================

def debug_print_table_structure(doc: Document) -> None:
    for t_idx, table in enumerate(doc.tables):
        print(f"\nTABLE {t_idx}: rows={len(table.rows)}, cols={len(table.columns)}")
        for r_idx, row in enumerate(table.rows):
            row_texts = []
            for c_idx, cell in enumerate(row.cells):
                row_texts.append(f"[{r_idx},{c_idx}] {cell.text[:80]!r}")
            print(" | ".join(row_texts))


def find_cell_with_label(table, label_text, partial=False):
    """
    Find a cell in the table containing the specified label text.
    Returns (cell, row_idx, col_idx) or (None, None, None) if not found.
    """
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            if partial:
                if label_text.lower() in cell_text.lower():
                    return cell, row_idx, col_idx
            else:
                if cell_text == label_text:
                    return cell, row_idx, col_idx
    return None, None, None


def fill_section_with_label(table, label_text, content):
    """
    Find a label and append content to that cell.
    Preserves the label text and appends content below it.
    """
    if not content or not safe_text(content):
        return False
    
    label_cell, row_idx, col_idx = find_cell_with_label(table, label_text, partial=True)
    
    if not label_cell:
        if DEBUG:
            print(f"  Label '{label_text}' not found")
        return False
    
    content_str = safe_text(content)
    cell_text = label_cell.text.strip()
    
    # If cell ends with ":" it's a label, append below it
    if cell_text.endswith(":"):
        label_cell.text = f"{cell_text}\n{content_str}"
    else:
        # Replace template text or append
        label_cell.text = content_str
    
    return True


def fill_template(template_path: str, row: dict, output_file: Path) -> None:
    """
    Fill template using label-based mapping.
    Searches for label text in cells instead of relying on hard-coded row numbers.
    """
    doc = Document(template_path)

    if len(doc.tables) == 0:
        raise ValueError("Template has no tables")

    table = doc.tables[0]

    if DEBUG:
        debug_print_table_structure(doc)

    # ========== PERSONAL DETAILS SECTION ==========
    fill_section_with_label(table, "Full Name", safe_text(row.get('FullName')))
    fill_section_with_label(table, "Advisor Name", safe_text(row.get('AdvisorName')))
    fill_section_with_label(table, "Address & Postcode", combine_address(row))
    fill_section_with_label(table, "Advisors contact", safe_text(row.get('CurrentSite')))
    fill_section_with_label(table, "Mobile Number", safe_text(row.get('MobileNumber')))
    fill_section_with_label(table, "Email Address", safe_text(row.get('Email')))

    # ========== BACKGROUND SECTION ==========
    background_text = build_background_text(row)
    fill_section_with_label(table, "Background", background_text)

    # ========== HEALTH WEEK 1 SECTION ==========
    week1_text = build_health_week1_text(row)
    fill_section_with_label(table, "Health Week1", week1_text)

    # ========== HEALTH WEEK 12 SECTION ==========
    week12_text = build_health_week12_text(row)
    fill_section_with_label(table, "Health Week 12", week12_text)

    # ========== CURRENT SUPPORT SECTION ==========
    support_text = build_support_text(row)
    fill_section_with_label(table, "Current level of support", support_text)

    # ========== WEEK 12 SUPPORT SECTION ==========
    supported_text = build_supported_text(row)
    fill_section_with_label(table, "Week 12 to they feel more supported", supported_text)

    # ========== REFERRAL SECTION ==========
    referral_text = build_referral_text(row)
    fill_section_with_label(table, "Referral made to and why", referral_text)

    # ========== COMMENTS SECTION ==========
    comments_text = build_comments_text(row)
    fill_section_with_label(table, "Any other comments", comments_text)

    # ========== GOALS SECTION ==========
    goals = split_goals(row)
    fill_section_with_label(table, "Goal A", goals[0])
    fill_section_with_label(table, "Goal B", goals[1])
    fill_section_with_label(table, "Goal C", goals[2])

    doc.save(output_file)

# =========================================================
# MAIN
# =========================================================

def main() -> None:
    Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)

    conn = get_connection()
    try:
        rows = fetch_assessment_pairs(conn)
    finally:
        conn.close()

    if not rows:
        print("No assessment records found.")
        return

    generated = 0
    skipped = 0

    for row in rows:
        if not has_current_assessment(row):
            skipped += 1
            print(f"Skipped {safe_text(row.get('SaheliCardNumber'))}: no current assessment")
            continue

        saheli_card = safe_text(row.get("SaheliCardNumber")) or f"PID_{safe_text(row.get('ParticipantID'))}"
        output_file = Path(OUTPUT_DIR) / f"{safe_filename(saheli_card)}.docx"

        try:
            fill_template(TEMPLATE_PATH, row, output_file)
            generated += 1
            print(f"Generated: {output_file}")
        except Exception as ex:
            skipped += 1
            print(f"Failed for {saheli_card}: {ex}")

    print(f"\nDone. Generated: {generated}, Skipped/Failed: {skipped}")
    print(f"Files saved in:\n{OUTPUT_DIR}")


if __name__ == "__main__":
    main()