import re
import shutil
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

import pandas as pd
from docx import Document
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from striprtf.striprtf import rtf_to_text


DEFAULT_CV_FOLDER = (
    Path.home()
    / "Documents"
    / "Saheli Recruitment"
    / "Health and Lifestyle Coordinator"
    / "CVs"
)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".rtf"}

# Scores total 100. The patterns are based only on job-related requirements.
CRITERIA: list[dict[str, Any]] = [
    {
        "name": "Physical activity or sports qualification",
        "weight": 12,
        "essential": True,
        "minimum_hits": 1,
        "patterns": [
            r"(?:level\s*[1-6]|certificate|diploma|degree|bsc|ba|msc|nvq|cyq|cimspa|reps|qualification|qualified).{0,80}(?:sport|fitness|exercise|physical activity|personal train|gym instruct|sports coach)",
            r"(?:sport|fitness|exercise|physical activity|personal train|gym instruct|sports coach).{0,80}(?:level\s*[1-6]|certificate|diploma|degree|bsc|ba|msc|nvq|cyq|cimspa|reps|qualification|qualified)",
            r"\bexercise referral\b",
            r"\blevel\s*2\s+(?:gym|fitness|exercise)\b",
            r"\blevel\s*3\s+(?:personal trainer|exercise|fitness)\b",
        ],
    },
    {
        "name": "Community, health or social care experience",
        "weight": 8,
        "essential": False,
        "minimum_hits": 2,
        "patterns": [
            r"\bcommunity health\b", r"\bhealth and wellbeing\b",
            r"\bhealth coach\b", r"\bsocial prescrib(?:er|ing)\b",
            r"\bcommunity development\b", r"\bsocial care\b",
            r"\bsupport worker\b", r"\bpublic health\b",
            r"\bcounsell(?:ing|or)\b", r"\bcommunity support\b",
        ],
    },
    {
        "name": "Outreach, referrals and partnerships",
        "weight": 12,
        "essential": True,
        "minimum_hits": 3,
        "patterns": [
            r"\boutreach\b", r"\bcommunity engagement\b",
            r"\breferral pathways?\b", r"\breferrals?\b",
            r"\bsocial prescribing\b", r"\bGP practices?\b",
            r"\bprimary care\b", r"\bNHS\b", r"\bVCFSE\b",
            r"\bpartner organisations?\b", r"\bcommunity partnerships?\b",
            r"\bstakeholder engagement\b",
        ],
    },
    {
        "name": "Personalised care, triage and caseload support",
        "weight": 10,
        "essential": True,
        "minimum_hits": 3,
        "patterns": [
            r"\bperson[- ]centred\b", r"\bpersonalised care\b",
            r"\bholistic support\b", r"\btriage\b",
            r"\binitial assessments?\b", r"\blifestyle assessments?\b",
            r"\bneeds assessments?\b", r"\bcaseload\b",
            r"\bcase management\b", r"\bsupport plans?\b",
            r"\baction plans?\b", r"\bcare plans?\b",
            r"\bone[- ]to[- ]one\b", r"\bgoal setting\b",
        ],
    },
    {
        "name": "Diet, nutrition and behaviour change",
        "weight": 9,
        "essential": True,
        "minimum_hits": 2,
        "patterns": [
            r"\bnutrition\b", r"\bhealthy eating\b", r"\bdietary\b",
            r"\bdiet plans?\b", r"\bfood choices?\b",
            r"\bweight management\b", r"\bdiabetes\b",
            r"\bblood sugar\b", r"\blong[- ]term conditions?\b",
            r"\bSouth Asian diet\b", r"\blifestyle change\b",
            r"\bbehaviou?r change\b",
        ],
    },
    {
        "name": "Safeguarding children and vulnerable adults",
        "weight": 8,
        "essential": True,
        "minimum_hits": 1,
        "patterns": [
            r"\bsafeguarding\b", r"\bchild protection\b",
            r"\badult protection\b", r"\bvulnerable adults?\b",
        ],
    },
    {
        "name": "CRM, monitoring, evaluation and reporting",
        "weight": 10,
        "essential": True,
        "minimum_hits": 3,
        "patterns": [
            r"\bCRM\b", r"\bREDCap\b", r"\bclient database\b",
            r"\bcase management system\b", r"\bmonitoring and evaluation\b",
            r"\bfunder reporting\b", r"\bquarterly reports?\b",
            r"\bperformance targets?\b", r"\bKPIs?\b",
            r"\boutcome measures?\b", r"\bbaseline\b",
            r"\bfollow[- ]up assessments?\b", r"\bprogress notes?\b",
            r"\bcase stud(?:y|ies)\b", r"\brecord keeping\b",
            r"\bdata collection\b", r"\breports?\b",
        ],
    },
    {
        "name": "Diverse communities and health inequalities",
        "weight": 7,
        "essential": True,
        "minimum_hits": 2,
        "patterns": [
            r"\bethnically diverse\b", r"\bdiverse communities\b",
            r"\bculturally sensitive\b", r"\bcultural awareness\b",
            r"\bminority communities\b", r"\bmigrant communities\b",
            r"\brefugees?\b", r"\bhealth inequalities\b",
            r"\bdeprivation\b", r"\bbarriers to access\b",
            r"\bequal opportunities\b", r"\bseldom heard\b",
            r"\bunderserved communities\b",
        ],
    },
    {
        "name": "Volunteer recruitment, training and support",
        "weight": 6,
        "essential": True,
        "minimum_hits": 2,
        "patterns": [
            r"\brecruit(?:ed|ing)? volunteers?\b",
            r"\btrain(?:ed|ing)? volunteers?\b",
            r"\bsupervis(?:e|ed|ing) volunteers?\b",
            r"\bsupport(?:ed|ing)? volunteers?\b",
            r"\bmotivat(?:e|ed|ing) volunteers?\b",
            r"\bvolunteer coordinator\b", r"\bvolunteer management\b",
        ],
    },
    {
        "name": "Communication and stakeholder engagement",
        "weight": 6,
        "essential": True,
        "minimum_hits": 2,
        "patterns": [
            r"\bstakeholder engagement\b", r"\bstakeholder management\b",
            r"\bpartnership working\b", r"\bnetworking meetings?\b",
            r"\bmulti[- ]agency\b", r"\bpresentation skills?\b",
            r"\bdelivered presentations?\b", r"\bfacilitat(?:e|ed|ing) workshops?\b",
            r"\bexcellent communication\b", r"\bactive listening\b",
            r"\brelationship building\b",
        ],
    },
    {
        "name": "IT, database and administration",
        "weight": 5,
        "essential": True,
        "minimum_hits": 2,
        "patterns": [
            r"\bMicrosoft Office\b", r"\bMicrosoft Excel\b", r"\bExcel\b",
            r"\bPowerPoint\b", r"\bMicrosoft Word\b", r"\bdatabase\b",
            r"\bdata entry\b", r"\badministration\b",
            r"\badministrative\b", r"\bIT systems?\b",
        ],
    },
    {
        "name": "Autonomous, team and flexible delivery",
        "weight": 5,
        "essential": True,
        "minimum_hits": 2,
        "patterns": [
            r"\bevenings? and weekends?\b", r"\bweekend working\b",
            r"\bflexible working\b", r"\bwork(?:ed|ing)? autonomously\b",
            r"\bindependent working\b", r"\bmanage(?:d|ment of)? own workload\b",
            r"\bprioritis(?:e|ed|ing)\b", r"\bwork(?:ed|ing)? under pressure\b",
            r"\bteam working\b", r"\bmultidisciplinary team\b",
            r"\bmulti[- ]site\b", r"\bphysical activity delivery\b",
        ],
    },
    {
        "name": "Additional language ability",
        "weight": 1,
        "essential": False,
        "minimum_hits": 1,
        "patterns": [
            r"\bbilingual\b", r"\bmultilingual\b", r"\bfluent in\b",
            r"\bnative speaker\b", r"\blanguages?\s*:",
        ],
    },
    {
        "name": "English and Maths qualification",
        "weight": 1,
        "essential": False,
        "minimum_hits": 1,
        "patterns": [
            r"\bGCSE.{0,60}(?:English|Maths|Mathematics)\b",
            r"\b(?:English|Maths|Mathematics).{0,60}GCSE\b",
            r"\bfunctional skills?.{0,50}(?:English|Maths)\b",
        ],
    },
]

QUALIFICATION_CRITERION = "Physical activity or sports qualification"

STATIC_EVIDENCE_CRITERIA = {
    QUALIFICATION_CRITERION,
    "Additional language ability",
    "English and Maths qualification",
}

ACTION_PATTERN = re.compile(
    r"\b(?:"
    r"manag(?:e|es|ed|ing)|"
    r"deliver(?:s|ed|ing)?|"
    r"support(?:s|ed|ing)?|"
    r"complet(?:e|es|ed|ing)|"
    r"refer(?:s|red|ring)?|"
    r"record(?:s|ed|ing)?|"
    r"report(?:s|ed|ing)?|"
    r"train(?:s|ed|ing)?|"
    r"supervis(?:e|es|ed|ing)|"
    r"develop(?:s|ed|ing)?|"
    r"coordinat(?:e|es|ed|ing)|"
    r"lead|led|leading|"
    r"creat(?:e|es|ed|ing)|"
    r"implement(?:s|ed|ing)?|"
    r"facilitat(?:e|es|ed|ing)|"
    r"assess(?:es|ed|ing)?|"
    r"monitor(?:s|ed|ing)?|"
    r"recruit(?:s|ed|ing)?|"
    r"achiev(?:e|es|ed|ing)|"
    r"improv(?:e|es|ed|ing)|"
    r"organis(?:e|es|ed|ing)|"
    r"maintain(?:s|ed|ing)?|"
    r"handl(?:e|es|ed|ing)|"
    r"us(?:e|es|ed|ing)|"
    r"produc(?:e|es|ed|ing)|"
    r"design(?:s|ed|ing)?|"
    r"evaluat(?:e|es|ed|ing)"
    r")\b",
    re.IGNORECASE,
)

SECTION_HEADINGS = {
    "profile": re.compile(
        r"^(?:personal |professional )?"
        r"(?:profile|summary|statement|objective|about me)\s*:?\s*$",
        re.IGNORECASE,
    ),
    "employment": re.compile(
        r"^(?:employment|work|professional|career)"
        r"(?: history| experience)?\s*:?\s*$",
        re.IGNORECASE,
    ),
    "education": re.compile(
        r"^(?:education(?:\s*(?:and|&)\s*qualifications?)?|"
        r"qualifications?|certifications?|"
        r"training|professional development)\s*:?\s*$",
        re.IGNORECASE,
    ),
    "skills": re.compile(
        r"^(?:key |core |professional )?"
        r"(?:skills|competencies|strengths)\s*:?\s*$",
        re.IGNORECASE,
    ),
}

DATE_PATTERN = re.compile(
    r"\b(?:(?:19|20)\d{2}|present|current)\b",
    re.IGNORECASE,
)

QUALIFICATION_PATTERNS = [
    re.compile(
        r"\blevel\s*2\b.{0,50}"
        r"\b(?:fitness|gym|exercise)\s+instructor\b",
        re.IGNORECASE,
    ),
    re.compile(
        r"\blevel\s*3\b.{0,50}"
        r"\b(?:personal train(?:er|ing)|fitness|exercise)\b",
        re.IGNORECASE,
    ),
    re.compile(
        r"\b(?:sports?|physical activity)\s+coach(?:ing)?\b"
        r".{0,60}\b(?:qualification|certificate|award|diploma|level)\b",
        re.IGNORECASE,
    ),
    re.compile(
        r"\b(?:qualification|certificate|award|diploma|level)\b"
        r".{0,60}\b(?:sports?|physical activity)\s+coach(?:ing)?\b",
        re.IGNORECASE,
    ),
    re.compile(
        r"\bexercise referral\b.{0,60}"
        r"\b(?:qualification|certificate|award|diploma|level)\b",
        re.IGNORECASE,
    ),
    re.compile(
        r"\b(?:qualification|certificate|award|diploma|level)\b"
        r".{0,60}\bexercise referral\b",
        re.IGNORECASE,
    ),
    re.compile(
        r"\b(?:degree|diploma|bsc|ba|msc)\b.{0,70}"
        r"\b(?:sport|sports science|physical activity|exercise science)\b",
        re.IGNORECASE,
    ),
    re.compile(
        r"\b(?:sport|sports science|physical activity|exercise science)\b"
        r".{0,70}\b(?:degree|diploma|bsc|ba|msc)\b",
        re.IGNORECASE,
    ),
]

JOB_DESCRIPTION_PHRASES = [
    "physical activity qualification",
    "community health",
    "health and wellbeing",
    "community engagement",
    "referral pathway",
    "social prescribing",
    "partner organisations",
    "stakeholder engagement",
    "personalised care",
    "initial assessment",
    "lifestyle assessment",
    "case management",
    "support plan",
    "goal setting",
    "healthy eating",
    "weight management",
    "behaviour change",
    "vulnerable adults",
    "monitoring and evaluation",
    "funder reporting",
    "performance targets",
    "outcome measures",
    "health inequalities",
    "barriers to access",
    "volunteer management",
    "partnership working",
    "active listening",
    "relationship building",
    "data entry",
    "flexible working",
    "work autonomously",
]

VERIFICATION_QUESTIONS = {
    QUALIFICATION_CRITERION: [
        (
            "What physical activity qualification do you hold? "
            "Please provide the certificate."
        ),
    ],
    "Outreach, referrals and partnerships": [
        "Please describe an outreach activity you personally delivered.",
        (
            "How many referrals did you generate, and which organisations "
            "referred people?"
        ),
    ],
    "CRM, monitoring, evaluation and reporting": [
        "Which CRM or case-management system did you use?",
        (
            "How did you measure whether a client's health or wellbeing "
            "improved?"
        ),
    ],
    "Safeguarding children and vulnerable adults": [
        (
            "Describe a safeguarding concern you handled and the "
            "procedure you followed."
        ),
    ],
    "Diet, nutrition and behaviour change": [
        "Give an example of a lifestyle or dietary plan you created.",
    ],
}

GROUP_STRONG = "Strong documented evidence"
GROUP_POSSIBLE = "Possible match - verification required"
GROUP_LIMITED = "Keyword match but limited supporting evidence"


def ask_cv_folder() -> Path:
    print(f"Default CV folder:\n{DEFAULT_CV_FOLDER}\n")
    value = input("Press ENTER to use it, or paste another CV folder path: ").strip().strip('"')
    return Path(value).expanduser() if value else DEFAULT_CV_FOLDER


def ask_shortlist_count() -> int:
    while True:
        value = input("How many candidates should be copied into the shortlist pack? [20]: ").strip()
        if not value:
            return 20
        try:
            count = int(value)
            if 1 <= count <= 500:
                return count
        except ValueError:
            pass
        print("Enter a whole number between 1 and 500.")


def extract_text(file_path: Path) -> tuple[str, str]:
    try:
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            reader = PdfReader(str(file_path))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
        elif suffix == ".docx":
            document = Document(str(file_path))
            parts = [p.text for p in document.paragraphs if p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if values:
                        parts.append(" | ".join(values))
            text = "\n".join(parts)
        elif suffix == ".txt":
            text = file_path.read_text(encoding="utf-8", errors="ignore")
        elif suffix == ".rtf":
            raw = file_path.read_text(encoding="utf-8", errors="ignore")
            text = rtf_to_text(raw)
        else:
            return "", f"Unsupported file type: {suffix}"

        text = text.strip()
        if len(text) < 150:
            return text, "Very little text was extracted. The CV may be scanned and needs manual review."
        return text, ""
    except Exception as error:
        return "", f"Could not extract text: {error}"


def normalise(text: str) -> str:
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\u00a0", " ").replace("–", "-").replace("—", "-")
    return re.sub(r"\s+", " ", text).lower().strip()


def shorten(text: str, maximum_length: int = 240) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) <= maximum_length:
        return text
    return text[: maximum_length - 3].rstrip() + "..."


def section_heading(line: str) -> str:
    for section, pattern in SECTION_HEADINGS.items():
        if pattern.fullmatch(line.strip()):
            return section
    return ""


def build_evidence_units(text: str) -> list[dict[str, Any]]:
    units: list[dict[str, Any]] = []
    current_section = "profile"

    for raw_line in text.splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip(" \t|:-")
        if not line:
            continue

        heading = section_heading(line)
        if heading:
            current_section = heading
            continue

        if DATE_PATTERN.search(line) and current_section != "education":
            current_section = "employment"
            line_section = "employment"
        else:
            line_section = current_section

        fragments = re.split(r"(?<=[.!?])\s+", line)
        for fragment in fragments:
            fragment = re.sub(r"\s+", " ", fragment).strip()
            if not fragment:
                continue
            units.append(
                {
                    "text": fragment,
                    "normalised": normalise(fragment),
                    "section": line_section,
                }
            )

    return units


def action_near_pattern(
    compiled_pattern: re.Pattern,
    text: str,
    maximum_distance: int = 160,
) -> bool:
    keyword_matches = list(compiled_pattern.finditer(text))
    action_matches = list(ACTION_PATTERN.finditer(text))

    for keyword_match in keyword_matches:
        keyword_centre = (
            keyword_match.start() + keyword_match.end()
        ) // 2

        for action_match in action_matches:
            action_centre = (
                action_match.start() + action_match.end()
            ) // 2
            if abs(keyword_centre - action_centre) <= maximum_distance:
                return True

    return False


def evidence_multiplier(
    section: str,
    action_required: bool,
) -> float:
    if action_required:
        return {
            "employment": 1.0,
            "other": 0.65,
            "education": 0.55,
            "profile": 0.35,
            "skills": 0.25,
        }.get(section, 0.50)

    return {
        "education": 1.0,
        "employment": 0.90,
        "other": 0.70,
        "skills": 0.60,
        "profile": 0.45,
    }.get(section, 0.60)


def find_specific_qualification(
    units: list[dict[str, Any]],
) -> tuple[str, float]:
    matches: list[tuple[float, str]] = []

    for unit in units:
        for pattern in QUALIFICATION_PATTERNS:
            if pattern.search(unit["text"]):
                quality = evidence_multiplier(
                    unit["section"],
                    action_required=False,
                )
                matches.append(
                    (
                        quality,
                        (
                            f"[{unit['section'].title()}] "
                            f"{shorten(unit['text'])}"
                        ),
                    )
                )
                break

    if not matches:
        return "", 0.0

    quality, evidence = max(matches, key=lambda item: item[0])
    return evidence, quality


def score_criterion(
    criterion: dict[str, Any],
    clean_text: str,
    units: list[dict[str, Any]],
    qualification_evidence: str,
    qualification_quality: float,
) -> dict[str, Any]:
    name = criterion["name"]
    weight = float(criterion["weight"])
    minimum_hits = max(int(criterion["minimum_hits"]), 1)
    action_required = name not in STATIC_EVIDENCE_CRITERIA

    compiled_patterns = [
        re.compile(pattern, re.IGNORECASE)
        for pattern in criterion["patterns"]
    ]
    matched_pattern_indexes = [
        index
        for index, pattern in enumerate(compiled_patterns)
        if pattern.search(clean_text)
    ]

    keyword_coverage = min(
        len(matched_pattern_indexes) / minimum_hits,
        1.0,
    )
    keyword_score = round(weight * keyword_coverage, 2)

    if name == QUALIFICATION_CRITERION:
        evidence_score = round(
            weight * qualification_quality,
            2,
        )
        return {
            "keyword_score": keyword_score,
            "evidence_score": evidence_score,
            "keyword_hits": len(matched_pattern_indexes),
            "evidence_count": 1 if qualification_evidence else 0,
            "evidence": qualification_evidence,
            "met": bool(qualification_evidence),
            "best_quality": qualification_quality,
        }

    candidates: list[tuple[float, int, int, str]] = []

    for pattern_index in matched_pattern_indexes:
        compiled_pattern = compiled_patterns[pattern_index]

        for unit_index, unit in enumerate(units):
            if not compiled_pattern.search(unit["text"]):
                continue

            has_context = (
                not action_required
                or action_near_pattern(
                    compiled_pattern,
                    unit["text"],
                )
            )
            if not has_context:
                continue

            quality = evidence_multiplier(
                unit["section"],
                action_required,
            )
            evidence = (
                f"[{unit['section'].title()}] "
                f"{shorten(unit['text'])}"
            )
            candidates.append(
                (
                    quality,
                    pattern_index,
                    unit_index,
                    evidence,
                )
            )

    selected: list[tuple[float, int, int, str]] = []
    used_patterns: set[int] = set()
    used_units: set[int] = set()

    for candidate in sorted(
        candidates,
        key=lambda item: item[0],
        reverse=True,
    ):
        _, pattern_index, unit_index, _ = candidate

        if pattern_index in used_patterns or unit_index in used_units:
            continue

        selected.append(candidate)
        used_patterns.add(pattern_index)
        used_units.add(unit_index)

        if len(selected) >= minimum_hits:
            break

    evidence_coverage = min(
        sum(item[0] for item in selected) / minimum_hits,
        1.0,
    )
    evidence_score = round(weight * evidence_coverage, 2)
    best_quality = max(
        (item[0] for item in selected),
        default=0.0,
    )

    met = (
        keyword_coverage >= 1.0
        and bool(selected)
        and best_quality >= 0.65
    )

    return {
        "keyword_score": keyword_score,
        "evidence_score": evidence_score,
        "keyword_hits": len(matched_pattern_indexes),
        "evidence_count": len(selected),
        "evidence": " | ".join(item[3] for item in selected[:2]),
        "met": met,
        "best_quality": best_quality,
    }


def job_description_overlap(
    clean_text: str,
) -> tuple[float, str, str]:
    matched_phrases = [
        phrase
        for phrase in JOB_DESCRIPTION_PHRASES
        if phrase in clean_text
    ]
    overlap = (
        len(matched_phrases) / len(JOB_DESCRIPTION_PHRASES)
        if JOB_DESCRIPTION_PHRASES
        else 0.0
    )
    warning = ""

    if len(matched_phrases) >= 8 and overlap >= 0.30:
        warning = (
            "High wording overlap with job description - "
            "manual verification required"
        )

    return overlap, warning, "; ".join(matched_phrases)


def verification_questions(
    criterion_results: dict[str, dict[str, Any]],
    overlap_warning: str,
) -> list[str]:
    questions: list[str] = []

    def add_questions(criterion_name: str) -> None:
        for question in VERIFICATION_QUESTIONS.get(
            criterion_name,
            [],
        ):
            if question not in questions:
                questions.append(question)

    add_questions(QUALIFICATION_CRITERION)

    ratios = {
        criterion["name"]: (
            criterion_results[criterion["name"]]["evidence_score"]
            / float(criterion["weight"])
        )
        for criterion in CRITERIA
    }

    for criterion_name in VERIFICATION_QUESTIONS:
        if criterion_name == QUALIFICATION_CRITERION:
            continue
        if ratios.get(criterion_name, 0.0) < 0.75:
            add_questions(criterion_name)

    if len(questions) < 4:
        for criterion_name, _ in sorted(
            ratios.items(),
            key=lambda item: item[1],
        ):
            add_questions(criterion_name)
            if len(questions) >= 4:
                break

    if overlap_warning:
        questions.append(
            (
                "Your CV closely follows the job-description wording. "
                "Please identify which examples are your own work and "
                "provide dates and a referee who can verify them."
            )
        )

    return questions


def assign_review_group(
    evidence_score: float,
    qualification_stated: bool,
    essential_met: int,
    essential_total: int,
) -> str:
    essential_ratio = (
        essential_met / essential_total
        if essential_total
        else 0.0
    )

    if (
        evidence_score >= 65.0
        and qualification_stated
        and essential_ratio >= 0.70
    ):
        return GROUP_STRONG

    if evidence_score >= 35.0:
        return GROUP_POSSIBLE

    return GROUP_LIMITED


def score_cv(file_path: Path) -> dict[str, Any]:
    text, warning = extract_text(file_path)
    clean_text = normalise(text)
    units = build_evidence_units(text)
    qualification_evidence, qualification_quality = (
        find_specific_qualification(units)
    )
    essential_total = sum(
        1
        for criterion in CRITERIA
        if criterion["essential"]
    )

    result: dict[str, Any] = {
        "Candidate": file_path.stem,
        "CV File": file_path.name,
        "Original CV Path": str(file_path.resolve()),
        "Evidence Quality Score": 0.0,
        "Keyword Match Score": 0.0,
        "Review Group": GROUP_LIMITED,
        "Essential Criteria Met": 0,
        "Essential Criteria Total": essential_total,
        "Specific Qualification Stated": (
            "Yes" if qualification_evidence else "No"
        ),
        "Qualification Status": (
            "Specific qualification stated - certificate required"
            if qualification_evidence
            else "No specific qualification found"
        ),
        "Qualification Evidence": qualification_evidence,
        "Essential Gaps": "",
        "Strong Evidence": "",
        "Job Description Wording Overlap": 0.0,
        "Wording Overlap Warning": "",
        "Matched Job Description Phrases": "",
        "Verification Questions": "",
        "Manual Review Notes": warning,
    }

    keyword_total = 0.0
    evidence_total = 0.0
    essential_met = 0
    essential_gaps: list[str] = []
    strengths: list[tuple[float, int, str]] = []
    criterion_results: dict[str, dict[str, Any]] = {}

    for criterion in CRITERIA:
        criterion_result = score_criterion(
            criterion,
            clean_text,
            units,
            qualification_evidence,
            qualification_quality,
        )
        criterion_results[criterion["name"]] = criterion_result

        keyword_total += criterion_result["keyword_score"]
        evidence_total += criterion_result["evidence_score"]

        result[
            f"Keyword Score - {criterion['name']}"
        ] = criterion_result["keyword_score"]
        result[
            f"Evidence Score - {criterion['name']}"
        ] = criterion_result["evidence_score"]
        result[
            f"Distinct Keyword Hits - {criterion['name']}"
        ] = criterion_result["keyword_hits"]
        result[
            f"Contextual Evidence Count - {criterion['name']}"
        ] = criterion_result["evidence_count"]
        result[
            f"Evidence - {criterion['name']}"
        ] = criterion_result["evidence"]

        if criterion["essential"]:
            if criterion_result["met"]:
                essential_met += 1
            else:
                essential_gaps.append(criterion["name"])

        if criterion_result["evidence"]:
            evidence_ratio = (
                criterion_result["evidence_score"]
                / float(criterion["weight"])
            )
            strengths.append(
                (
                    evidence_ratio,
                    int(criterion["weight"]),
                    (
                        f"{criterion['name']}: "
                        f"{criterion_result['evidence']}"
                    ),
                )
            )

    overlap, overlap_warning, matched_phrases = (
        job_description_overlap(clean_text)
    )
    questions = verification_questions(
        criterion_results,
        overlap_warning,
    )

    result["Keyword Match Score"] = (
        round(keyword_total, 2)
        if text
        else 0.0
    )
    result["Evidence Quality Score"] = (
        round(evidence_total, 2)
        if text
        else 0.0
    )
    result["Essential Criteria Met"] = essential_met
    result["Essential Gaps"] = "; ".join(essential_gaps)
    result["Strong Evidence"] = " || ".join(
        item[2]
        for item in sorted(
            strengths,
            reverse=True,
        )[:4]
    )
    result["Job Description Wording Overlap"] = overlap
    result["Wording Overlap Warning"] = overlap_warning
    result["Matched Job Description Phrases"] = matched_phrases
    result["Verification Questions"] = "\n".join(
        f"{number}. {question}"
        for number, question in enumerate(
            questions,
            start=1,
        )
    )
    result["Review Group"] = assign_review_group(
        result["Evidence Quality Score"],
        bool(qualification_evidence),
        essential_met,
        essential_total,
    )

    review_notes = [
        str(result["Manual Review Notes"] or "").strip(),
        overlap_warning,
    ]
    if not qualification_evidence:
        review_notes.append(
            (
                "No specific physical activity qualification was "
                "identified - verify certificates manually"
            )
        )
    result["Manual Review Notes"] = "; ".join(
        note
        for note in review_notes
        if note
    )

    if not text and not warning:
        result["Manual Review Notes"] = (
            "No readable CV text was found."
        )

    return result


def format_sheet(worksheet) -> None:
    worksheet.freeze_panes = "A2"
    worksheet.auto_filter.ref = worksheet.dimensions
    worksheet.sheet_view.showGridLines = False

    header_fill = PatternFill(
        fill_type="solid",
        fgColor="1F4E3D",
    )
    thin_border = Border(
        bottom=Side(
            style="thin",
            color="D5E3DC",
        )
    )

    for cell in worksheet[1]:
        cell.font = Font(
            bold=True,
            color="FFFFFF",
        )
        cell.fill = header_fill
        cell.alignment = Alignment(
            horizontal="center",
            vertical="center",
            wrap_text=True,
        )

    worksheet.row_dimensions[1].height = 32

    for row in worksheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(
                vertical="top",
                wrap_text=True,
            )
            cell.border = thin_border

    headers = {
        cell.value: cell.column
        for cell in worksheet[1]
    }

    for heading in [
        "Evidence Quality Score",
        "Keyword Match Score",
    ]:
        column = headers.get(heading)
        if column:
            for row_number in range(2, worksheet.max_row + 1):
                worksheet.cell(
                    row=row_number,
                    column=column,
                ).number_format = "0.0"

    overlap_column = headers.get(
        "Job Description Wording Overlap"
    )
    if overlap_column:
        for row_number in range(2, worksheet.max_row + 1):
            worksheet.cell(
                row=row_number,
                column=overlap_column,
            ).number_format = "0%"

    group_column = headers.get("Review Group")
    if group_column:
        group_fills = {
            GROUP_STRONG: PatternFill(
                fill_type="solid",
                fgColor="D9EAD3",
            ),
            GROUP_POSSIBLE: PatternFill(
                fill_type="solid",
                fgColor="FFF2CC",
            ),
            GROUP_LIMITED: PatternFill(
                fill_type="solid",
                fgColor="FCE8E6",
            ),
        }
        for row_number in range(2, worksheet.max_row + 1):
            cell = worksheet.cell(
                row=row_number,
                column=group_column,
            )
            if cell.value in group_fills:
                cell.fill = group_fills[cell.value]
                cell.font = Font(bold=True)

    warning_column = headers.get("Wording Overlap Warning")
    if warning_column:
        warning_fill = PatternFill(
            fill_type="solid",
            fgColor="FCE8E6",
        )
        for row_number in range(2, worksheet.max_row + 1):
            cell = worksheet.cell(
                row=row_number,
                column=warning_column,
            )
            if str(cell.value or "").strip():
                cell.fill = warning_fill
                cell.font = Font(
                    bold=True,
                    color="9C0006",
                )

    for column_cells in worksheet.columns:
        letter = column_cells[0].column_letter
        header = str(column_cells[0].value or "")
        longest = max(
            len(str(cell.value or ""))
            for cell in column_cells
        )

        if any(
            phrase in header
            for phrase in [
                "Evidence",
                "Questions",
                "Notes",
                "Gaps",
                "Phrases",
                "Rule",
                "Guidance",
            ]
        ):
            maximum_width = 60
        elif header in {
            "Candidate",
            "Review Group",
            "Qualification Status",
        }:
            maximum_width = 34
        else:
            maximum_width = 22

        worksheet.column_dimensions[letter].width = min(
            max(longest + 2, 12),
            maximum_width,
        )

    for row_number in range(2, worksheet.max_row + 1):
        maximum_lines = 1

        for cell in worksheet[row_number]:
            text = str(cell.value or "")
            column_letter = cell.column_letter
            width = int(
                worksheet.column_dimensions[
                    column_letter
                ].width
                or 12
            )
            estimated_lines = sum(
                max(
                    1,
                    (len(line) + width - 1) // width,
                )
                for line in text.splitlines()
            )
            maximum_lines = max(
                maximum_lines,
                estimated_lines,
            )

        worksheet.row_dimensions[row_number].height = min(
            max(18, maximum_lines * 14),
            300,
        )


def pdf_text(
    value: Any,
    empty_text: str = "Not provided",
) -> str:
    if value is None:
        return empty_text

    try:
        if pd.isna(value):
            return empty_text
    except (TypeError, ValueError):
        pass

    text = str(value).strip()
    if not text:
        return empty_text

    text = "".join(
        character
        if character in "\n\t" or ord(character) >= 32
        else " "
        for character in text
    )
    return escape(text).replace("\n", "<br/>")


def add_pdf_footer(canvas, document) -> None:
    canvas.saveState()
    canvas.setStrokeColor(colors.HexColor("#D9EAD3"))
    canvas.line(
        document.leftMargin,
        13 * mm,
        A4[0] - document.rightMargin,
        13 * mm,
    )
    canvas.setFillColor(colors.HexColor("#5F6368"))
    canvas.setFont("Helvetica", 8)
    canvas.drawString(
        document.leftMargin,
        8 * mm,
        "Health and Lifestyle Coordinator CV Shortlist",
    )
    canvas.drawRightString(
        A4[0] - document.rightMargin,
        8 * mm,
        f"Page {canvas.getPageNumber()}",
    )
    canvas.restoreState()


def create_shortlist_pdf(
    top: pd.DataFrame,
    output_path: Path,
) -> None:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ShortlistTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=24,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#1F4E3D"),
        spaceAfter=5 * mm,
    )
    subtitle_style = ParagraphStyle(
        "ShortlistSubtitle",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#5F6368"),
        spaceAfter=6 * mm,
    )
    heading_style = ParagraphStyle(
        "CandidateHeading",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#1F4E3D"),
        spaceAfter=4 * mm,
    )
    section_style = ParagraphStyle(
        "CandidateSection",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=13,
        textColor=colors.HexColor("#1F4E3D"),
        spaceBefore=4 * mm,
        spaceAfter=1.5 * mm,
    )
    body_style = ParagraphStyle(
        "ShortlistBody",
        parent=styles["BodyText"],
        fontSize=9,
        leading=13,
        textColor=colors.HexColor("#202124"),
        spaceAfter=2 * mm,
    )
    warning_style = ParagraphStyle(
        "ShortlistWarning",
        parent=body_style,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#9C0006"),
    )
    table_header_style = ParagraphStyle(
        "ShortlistTableHeader",
        parent=body_style,
        fontName="Helvetica-Bold",
        textColor=colors.white,
        alignment=TA_CENTER,
    )
    label_style = ParagraphStyle(
        "ShortlistLabel",
        parent=body_style,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#1F4E3D"),
    )

    document = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=16 * mm,
        bottomMargin=20 * mm,
        title="Health and Lifestyle Coordinator CV Shortlist",
        author="Saheli Hub",
    )

    story = [
        Paragraph(
            "Health and Lifestyle Coordinator",
            title_style,
        ),
        Paragraph("CV Shortlist", title_style),
        Paragraph(
            (
                f"Top {len(top)} candidates - "
                f"created {datetime.now():%d %B %Y at %H:%M}"
            ),
            subtitle_style,
        ),
    ]

    group_counts = top["Review Group"].value_counts()
    story.extend(
        [
            Paragraph(
                (
                    f"<b>{GROUP_STRONG}:</b> "
                    f"{int(group_counts.get(GROUP_STRONG, 0))}"
                ),
                body_style,
            ),
            Paragraph(
                (
                    f"<b>{GROUP_POSSIBLE}:</b> "
                    f"{int(group_counts.get(GROUP_POSSIBLE, 0))}"
                ),
                body_style,
            ),
            Paragraph(
                (
                    f"<b>{GROUP_LIMITED}:</b> "
                    f"{int(group_counts.get(GROUP_LIMITED, 0))}"
                ),
                body_style,
            ),
            Spacer(1, 3 * mm),
        ]
    )

    summary_data = [
        [
            Paragraph("Rank", table_header_style),
            Paragraph("Candidate", table_header_style),
            Paragraph("Review group", table_header_style),
            Paragraph("Evidence", table_header_style),
            Paragraph("Keyword", table_header_style),
            Paragraph("Qual.", table_header_style),
        ]
    ]

    for _, candidate in top.iterrows():
        summary_data.append(
            [
                Paragraph(
                    str(int(candidate["Review Rank"])),
                    body_style,
                ),
                Paragraph(
                    pdf_text(candidate["Candidate"]),
                    body_style,
                ),
                Paragraph(
                    pdf_text(candidate["Review Group"]),
                    body_style,
                ),
                Paragraph(
                    f"{float(candidate['Evidence Quality Score']):.1f}",
                    body_style,
                ),
                Paragraph(
                    f"{float(candidate['Keyword Match Score']):.1f}",
                    body_style,
                ),
                Paragraph(
                    pdf_text(
                        candidate["Specific Qualification Stated"]
                    ),
                    body_style,
                ),
            ]
        )

    summary_table = Table(
        summary_data,
        colWidths=[
            12 * mm,
            53 * mm,
            47 * mm,
            22 * mm,
            22 * mm,
            22 * mm,
        ],
        repeatRows=1,
        hAlign="LEFT",
    )
    summary_table.setStyle(
        TableStyle(
            [
                (
                    "BACKGROUND",
                    (0, 0),
                    (-1, 0),
                    colors.HexColor("#1F4E3D"),
                ),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                (
                    "GRID",
                    (0, 0),
                    (-1, -1),
                    0.4,
                    colors.HexColor("#C9D8D0"),
                ),
                (
                    "ROWBACKGROUNDS",
                    (0, 1),
                    (-1, -1),
                    [
                        colors.white,
                        colors.HexColor("#F4F8F6"),
                    ],
                ),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.extend(
        [
            summary_table,
            Spacer(1, 5 * mm),
            Paragraph(
                (
                    "Review order is driven primarily by documented "
                    "evidence quality. Keyword matches alone do not "
                    "confirm competence or qualifications."
                ),
                body_style,
            ),
        ]
    )

    for _, candidate in top.iterrows():
        story.append(PageBreak())
        story.append(
            Paragraph(
                (
                    f"#{int(candidate['Review Rank'])} - "
                    f"{pdf_text(candidate['Candidate'])}"
                ),
                heading_style,
            )
        )

        essential_result = (
            f"{int(candidate['Essential Criteria Met'])} of "
            f"{int(candidate['Essential Criteria Total'])} met"
        )
        details = [
            [
                Paragraph("Review group", label_style),
                Paragraph(
                    pdf_text(candidate["Review Group"]),
                    body_style,
                ),
            ],
            [
                Paragraph("Evidence quality score", label_style),
                Paragraph(
                    (
                        f"{float(candidate['Evidence Quality Score']):.1f} "
                        "out of 100"
                    ),
                    body_style,
                ),
            ],
            [
                Paragraph("Keyword match score", label_style),
                Paragraph(
                    (
                        f"{float(candidate['Keyword Match Score']):.1f} "
                        "out of 100"
                    ),
                    body_style,
                ),
            ],
            [
                Paragraph("Essential criteria", label_style),
                Paragraph(essential_result, body_style),
            ],
            [
                Paragraph("Qualification", label_style),
                Paragraph(
                    pdf_text(candidate["Qualification Status"]),
                    body_style,
                ),
            ],
            [
                Paragraph("Original CV", label_style),
                Paragraph(
                    pdf_text(candidate["CV File"]),
                    body_style,
                ),
            ],
        ]

        details_table = Table(
            details,
            colWidths=[
                44 * mm,
                134 * mm,
            ],
            hAlign="LEFT",
        )
        details_table.setStyle(
            TableStyle(
                [
                    (
                        "BACKGROUND",
                        (0, 0),
                        (0, -1),
                        colors.HexColor("#EAF2EE"),
                    ),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    (
                        "GRID",
                        (0, 0),
                        (-1, -1),
                        0.4,
                        colors.HexColor("#C9D8D0"),
                    ),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.append(details_table)

        story.append(
            Paragraph("Qualification evidence", section_style)
        )
        story.append(
            Paragraph(
                pdf_text(
                    candidate["Qualification Evidence"],
                    "No specific qualification evidence found.",
                ),
                body_style,
            )
        )

        story.append(
            Paragraph("Essential gaps", section_style)
        )
        story.append(
            Paragraph(
                pdf_text(
                    candidate["Essential Gaps"],
                    "No essential gaps detected by the script.",
                ),
                body_style,
            )
        )

        story.append(
            Paragraph("Strong documented evidence", section_style)
        )
        story.append(
            Paragraph(
                pdf_text(
                    candidate["Strong Evidence"],
                    "No contextual evidence extracted.",
                ).replace(" || ", "<br/><br/>"),
                body_style,
            )
        )

        overlap_warning = str(
            candidate["Wording Overlap Warning"] or ""
        ).strip()
        if overlap_warning:
            story.append(
                Paragraph(
                    "Wording-overlap warning",
                    section_style,
                )
            )
            story.append(
                Paragraph(
                    pdf_text(overlap_warning),
                    warning_style,
                )
            )

        story.append(
            Paragraph("Verification questions", section_style)
        )
        story.append(
            Paragraph(
                pdf_text(candidate["Verification Questions"]),
                body_style,
            )
        )

        review_notes = str(
            candidate["Manual Review Notes"] or ""
        ).strip()
        if review_notes:
            story.append(
                Paragraph("Manual review notes", section_style)
            )
            story.append(
                Paragraph(
                    pdf_text(review_notes),
                    warning_style,
                )
            )

    document.build(
        story,
        onFirstPage=add_pdf_footer,
        onLaterPages=add_pdf_footer,
    )


def create_pack(
    results: list[dict[str, Any]],
    shortlist_count: int,
    cv_folder: Path,
) -> tuple[Path, Path, Path, int]:
    dataframe = pd.DataFrame(results).sort_values(
        by=[
            "Evidence Quality Score",
            "Specific Qualification Stated",
            "Essential Criteria Met",
            "Keyword Match Score",
        ],
        ascending=[
            False,
            False,
            False,
            False,
        ],
        kind="stable",
    ).reset_index(drop=True)
    dataframe.insert(0, "Review Rank", range(1, len(dataframe) + 1))

    shortlist_count = min(shortlist_count, len(dataframe))
    top = dataframe.head(shortlist_count).copy()

    base_folder = cv_folder.parent
    pack_folder = base_folder / "Health_Lifestyle_Shortlist_Pack"
    shortlisted_cv_folder = pack_folder / "Shortlisted_CVs"
    excel_path = pack_folder / "Health_Lifestyle_CV_Shortlist.xlsx"
    csv_path = pack_folder / "Health_Lifestyle_All_Candidates.csv"
    pdf_path = pack_folder / "Health_Lifestyle_CV_Shortlist.pdf"
    zip_path = base_folder / "Health_Lifestyle_Shortlist_Pack.zip"

    if pack_folder.exists():
        shutil.rmtree(pack_folder)
    shortlisted_cv_folder.mkdir(parents=True, exist_ok=True)

    top["Shortlisted CV"] = ""
    copied_count = 0

    for index, row in top.iterrows():
        source = Path(row["Original CV Path"])
        if not source.exists():
            note = str(row.get("Manual Review Notes", "") or "").strip()
            extra = f"Original CV not found: {source}"
            top.at[index, "Manual Review Notes"] = f"{note}; {extra}" if note else extra
            continue

        destination_name = f"{int(row['Review Rank']):02d}_{source.name}"
        destination = shortlisted_cv_folder / destination_name
        shutil.copy2(source, destination)
        top.at[index, "Shortlisted CV"] = f"Shortlisted_CVs/{destination_name}"
        copied_count += 1

    shortlist_columns = [
        "Review Rank",
        "Review Group",
        "Candidate",
        "CV File",
        "Shortlisted CV",
        "Evidence Quality Score",
        "Keyword Match Score",
        "Essential Criteria Met",
        "Essential Criteria Total",
        "Specific Qualification Stated",
        "Qualification Status",
        "Qualification Evidence",
        "Essential Gaps",
        "Strong Evidence",
        "Job Description Wording Overlap",
        "Wording Overlap Warning",
        "Verification Questions",
        "Manual Review Notes",
    ]

    criteria_df = pd.DataFrame(
        [
            {
                "Criterion": criterion["name"],
                "Type": (
                    "Essential"
                    if criterion["essential"]
                    else "Desirable"
                ),
                "Maximum Weight": criterion["weight"],
                "Distinct Keyword Signals for Full Score": (
                    criterion["minimum_hits"]
                ),
                "Contextual Action Required": (
                    "No - specific credential or static fact"
                    if criterion["name"] in STATIC_EVIDENCE_CRITERIA
                    else "Yes - action must be near the keyword"
                ),
                "Evidence Rule": (
                    (
                        "Must state a specific sport, fitness, exercise "
                        "or physical-activity credential"
                    )
                    if criterion["name"] == QUALIFICATION_CRITERION
                    else (
                        "Each distinct keyword pattern and each evidence "
                        "statement can contribute only once"
                    )
                ),
            }
            for criterion in CRITERIA
        ]
    )

    group_guidance = {
        GROUP_STRONG: (
            "Evidence Quality Score at least 65, a specific physical "
            "activity qualification stated, and at least 70% of "
            "essential criteria supported. Verify certificates and "
            "references before appointment."
        ),
        GROUP_POSSIBLE: (
            "Some contextual evidence is present, but important claims "
            "need interview, certificate or reference verification."
        ),
        GROUP_LIMITED: (
            "Keywords are present but documented actions are limited. "
            "Review manually before deciding whether to interview."
        ),
    }
    group_summary_df = pd.DataFrame(
        [
            {
                "Review Group": group,
                "Candidate Count": int(
                    (dataframe["Review Group"] == group).sum()
                ),
                "Review Guidance": group_guidance[group],
            }
            for group in [
                GROUP_STRONG,
                GROUP_POSSIBLE,
                GROUP_LIMITED,
            ]
        ]
    )

    method_df = pd.DataFrame(
        [
            {
                "Method": "Review order",
                "Rule": (
                    "Rank primarily by Evidence Quality Score, then "
                    "specific qualification, essential criteria met, "
                    "and finally Keyword Match Score."
                ),
            },
            {
                "Method": "Keyword repetition cap",
                "Rule": (
                    "A distinct keyword pattern contributes at most once "
                    "per criterion, regardless of repetition."
                ),
            },
            {
                "Method": "Contextual evidence",
                "Rule": (
                    "Most criteria require an action word within 160 "
                    "characters of the matched term."
                ),
            },
            {
                "Method": "Section weighting",
                "Rule": (
                    "Dated employment evidence receives full weight; "
                    "profile claims receive 35% and skills-list claims "
                    "receive 25% when action context is required."
                ),
            },
            {
                "Method": "Qualification verification",
                "Rule": (
                    "The essential physical-activity qualification is "
                    "met only when a specific credential is stated. "
                    "The certificate must still be checked manually."
                ),
            },
            {
                "Method": "Wording overlap",
                "Rule": (
                    "High overlap creates a manual-verification warning "
                    "only; it does not reject or downgrade a candidate."
                ),
            },
        ]
    )

    question_rows: list[dict[str, Any]] = []
    for _, candidate in top.iterrows():
        for question_line in str(
            candidate["Verification Questions"] or ""
        ).splitlines():
            question = re.sub(
                r"^\d+\.\s*",
                "",
                question_line,
            ).strip()
            if question:
                question_rows.append(
                    {
                        "Review Rank": int(
                            candidate["Review Rank"]
                        ),
                        "Review Group": candidate["Review Group"],
                        "Candidate": candidate["Candidate"],
                        "Question": question,
                    }
                )

    questions_df = pd.DataFrame(
        question_rows,
        columns=[
            "Review Rank",
            "Review Group",
            "Candidate",
            "Question",
        ],
    )

    manual_review_df = dataframe[
        (
            dataframe["Manual Review Notes"]
            .fillna("")
            .astype(str)
            .str.strip()
            != ""
        )
    ][
        [
            "Review Rank",
            "Review Group",
            "Candidate",
            "CV File",
            "Evidence Quality Score",
            "Keyword Match Score",
            "Specific Qualification Stated",
            "Qualification Status",
            "Wording Overlap Warning",
            "Manual Review Notes",
            "Original CV Path",
        ]
    ].copy()

    dataframe.to_csv(csv_path, index=False, encoding="utf-8-sig")

    top_report = top[shortlist_columns].copy()
    top_report["Strong Evidence"] = top_report[
        "Strong Evidence"
    ].apply(
        lambda value: shorten(
            str(value or ""),
            maximum_length=520,
        )
    )
    top_report["Verification Questions"] = (
        "See the Verification Questions sheet"
    )

    with pd.ExcelWriter(excel_path, engine="openpyxl") as writer:
        top_report.to_excel(
            writer,
            sheet_name="Top Candidates",
            index=False,
        )
        group_summary_df.to_excel(
            writer,
            sheet_name="Review Groups",
            index=False,
        )
        questions_df.to_excel(
            writer,
            sheet_name="Verification Questions",
            index=False,
        )
        dataframe.to_excel(
            writer,
            sheet_name="All Candidates",
            index=False,
        )
        criteria_df.to_excel(
            writer,
            sheet_name="Scoring Criteria",
            index=False,
        )
        method_df.to_excel(
            writer,
            sheet_name="Scoring Method",
            index=False,
        )
        manual_review_df.to_excel(
            writer,
            sheet_name="Manual Review",
            index=False,
        )

        top_sheet = writer.book["Top Candidates"]
        cv_column = shortlist_columns.index("Shortlisted CV") + 1
        for excel_row, (_, candidate) in enumerate(top.iterrows(), start=2):
            link = str(candidate["Shortlisted CV"] or "").strip()
            cell = top_sheet.cell(row=excel_row, column=cv_column)
            if link:
                cell.value = "Open CV"
                cell.hyperlink = link
                cell.style = "Hyperlink"
            else:
                cell.value = "CV not copied"

        tab_colours = {
            "Top Candidates": "1F4E3D",
            "Review Groups": "5B9BD5",
            "Verification Questions": "ED7D31",
            "All Candidates": "70AD47",
            "Scoring Criteria": "A5A5A5",
            "Scoring Method": "A5A5A5",
            "Manual Review": "C00000",
        }
        for worksheet in writer.book.worksheets:
            worksheet.sheet_properties.tabColor = tab_colours.get(
                worksheet.title
            )
            format_sheet(worksheet)

    create_shortlist_pdf(top, pdf_path)

    if zip_path.exists():
        zip_path.unlink()
    shutil.make_archive(str(zip_path.with_suffix("")), "zip", root_dir=pack_folder)

    return pack_folder, zip_path, pdf_path, copied_count


def main() -> None:
    print("=" * 72)
    print("Health and Lifestyle Coordinator CV Shortlist Pack")
    print("=" * 72)
    print()

    cv_folder = ask_cv_folder()
    if not cv_folder.exists() or not cv_folder.is_dir():
        raise FileNotFoundError(f"CV folder was not found:\n{cv_folder}")

    shortlist_count = ask_shortlist_count()
    cv_files = sorted(
        [file for file in cv_folder.iterdir() if file.is_file() and file.suffix.lower() in SUPPORTED_EXTENSIONS],
        key=lambda file: file.name.lower(),
    )

    old_doc_files = sorted(cv_folder.glob("*.doc"))
    if old_doc_files:
        print("\nThese old .doc files need converting to PDF or DOCX before they can be scored:")
        for file in old_doc_files:
            print(f" - {file.name}")

    if not cv_files:
        raise RuntimeError("No readable CV files were found. Supported formats: PDF, DOCX, TXT and RTF.")

    print(f"\nCVs found: {len(cv_files)}")
    print(f"Shortlist requested: {shortlist_count}\n")

    results: list[dict[str, Any]] = []
    for number, cv_file in enumerate(cv_files, start=1):
        print(f"[{number}/{len(cv_files)}] Reviewing {cv_file.name}")
        results.append(score_cv(cv_file))

    pack_folder, zip_path, pdf_path, copied_count = create_pack(
        results,
        shortlist_count,
        cv_folder,
    )

    print("\n" + "=" * 72)
    print("Completed")
    print("=" * 72)
    print(f"Candidates reviewed: {len(results)}")
    print(f"Shortlisted CVs copied: {copied_count}")
    print(f"Pack folder: {pack_folder}")
    print(f"PDF shortlist report: {pdf_path}")
    print(f"ZIP file to share: {zip_path}")
    print(
        "\nThe ranking is a review aid. Check original CVs, "
        "certificates, references and interview evidence before deciding."
    )


if __name__ == "__main__":
    try:
        main()
    except PermissionError:
        print("\nERROR: Close the existing Excel/ZIP file and run the script again.")
    except Exception as error:
        print(f"\nERROR: {error}")
    input("\nPress ENTER to close...")
