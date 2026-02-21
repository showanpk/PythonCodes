from pathlib import Path
import sys


# Target input from your request.
INPUT_AUDIO = Path(r"C:\Users\shonk\OneDrive\Documents\Sound Recordings\Recording.m4a")
# Output .docx file in the same folder by default.
OUTPUT_DOCX = INPUT_AUDIO.with_suffix(".docx")
MODEL_NAME = "medium"


def resolve_input_path(raw_path: Path) -> Path:
    """
    Resolve common Windows + OneDrive path issues.
    """
    # Direct path check first.
    if raw_path.exists():
        return raw_path

    # Common fallback in case OneDrive root is different on this machine.
    fallback = (
        Path.home()
        / "OneDrive"
        / "Documents"
        / "Sound Recordings"
        / raw_path.name
    )
    if fallback.exists():
        return fallback

    raise FileNotFoundError(
        "Audio file not found.\n"
        f"Checked:\n- {raw_path}\n- {fallback}\n\n"
        "If this file is in OneDrive, make sure it is downloaded locally:\n"
        "Right-click the file in OneDrive > Always keep on this device."
    )


def save_transcript_to_docx(transcript_text: str, output_path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("Whisper Transcription", level=1)
    doc.add_paragraph(transcript_text.strip())
    doc.save(str(output_path))


def main() -> int:
    try:
        input_path = resolve_input_path(INPUT_AUDIO)
    except FileNotFoundError as err:
        print(f"\nERROR: {err}\n")
        return 1

    try:
        import whisper
    except ImportError:
        print(
            "ERROR: The 'whisper' package is not installed.\n"
            "Install with: pip install -U openai-whisper"
        )
        return 1

    try:
        import docx  # noqa: F401
    except ImportError:
        print(
            "ERROR: The 'python-docx' package is not installed.\n"
            "Install with: pip install python-docx"
        )
        return 1

    print(f"Loading Whisper model: {MODEL_NAME}")
    model = whisper.load_model(MODEL_NAME)

    print(f"Transcribing: {input_path}")
    print("Progress Bar:")
    # Whisper shows a tqdm progress bar when verbose=False.
    result = model.transcribe(str(input_path), verbose=False, fp16=False)

    transcript = result.get("text", "").strip()
    if not transcript:
        print("WARNING: No text was transcribed.")

    save_transcript_to_docx(transcript, OUTPUT_DOCX)
    print(f"\nDone. Transcript saved to:\n{OUTPUT_DOCX}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
